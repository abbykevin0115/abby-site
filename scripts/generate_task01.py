from __future__ import annotations

import argparse
import datetime as dt
import re
from collections import Counter
from pathlib import Path

import yaml
from docx import Document


# -----------------------
# 讀取模板（任務規格）
# -----------------------
def load_template(template_path: Path) -> dict:
    with template_path.open("r", encoding="utf-8") as f:
        return yaml.safe_load(f)


# -----------------------
# 讀取來源文字（30 天文章 / 全系列 Word）
# -----------------------
def read_docx_paragraphs(source_path: Path) -> list[str]:
    if source_path.suffix.lower() != ".docx":
        raise ValueError(f"目前只支援 docx，收到：{source_path.suffix}")
    doc = Document(str(source_path))
    return [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]


# -----------------------
# 判斷某段落是否「像程式碼」
# 目的：把 code-heavy 段落先排除，避免抽到 const/if/for 這種 token
# -----------------------
_CODE_SYMBOLS = set("{}[]();=<>/*\\|`~$#@")
_CODE_KEYWORDS = {
    "const", "let", "var", "function", "return", "if", "else", "for", "while",
    "class", "import", "export", "from", "try", "catch", "async", "await",
    "def", "print", "None", "True", "False"
}


def looks_like_code(paragraph: str) -> bool:
    s = paragraph.strip()
    if not s:
        return False

    # 1) 太短且只有符號/代碼
    if len(s) <= 25 and any(ch in _CODE_SYMBOLS for ch in s):
        return True

    # 2) 符號比例高（常見於 code）
    symbol_count = sum(1 for ch in s if ch in _CODE_SYMBOLS)
    if symbol_count / max(len(s), 1) >= 0.08:
        return True

    # 3) 出現多個 code keywords
    tokens = re.findall(r"[A-Za-z_]\w*", s)
    hit = sum(1 for t in tokens if t in _CODE_KEYWORDS)
    if hit >= 2:
        return True

    # 4) 像 JSON / dict / code block
    if s.startswith("{") and s.endswith("}"):
        return True
    if "```" in s:
        return True

    return False


# -----------------------
# 抽取「技術名詞候選」（通用版，偏出版可用）
# 策略：
# - 只從「非程式碼段落」抽（先過濾）
# - 優先抓：縮寫（LLM/RAG/API）、TitleCase（LangGraph/NotebookLM）、含連字號/點的框架名
# - 排除：過短/常見字/程式語法字
# -----------------------
_STOP_WORDS = {
    "the","and","with","from","into","that","this","your","you","are","for","to","in","of","on","as",
    "is","be","by","an","or","at","it","we","our","can","will","not","use","using",
    "a","i","ok","app"  # 你剛剛踩雷的也直接放進來
}
_BLACKLIST = {
    # 常見程式 token 或太泛用，不該進故事當「技術名詞」
    "const", "let", "var", "if", "else", "for", "while", "return",
    "true", "false", "none", "null",
    "amount", "category", "title", "data", "test", "example"
}


def extract_terms_from_text(text: str, top_k: int = 12) -> list[str]:
    # 1) 先抓英文字串候選（含 - _ .）
    raw = re.findall(r"[A-Za-z][A-Za-z0-9_\-\.]{1,}", text)

    candidates: list[str] = []
    for w in raw:
        lw = w.lower()

        if lw in _STOP_WORDS or lw in _BLACKLIST:
            continue
        if len(w) < 3:
            continue

        # 優先保留：縮寫（全大寫）或 TitleCase 或 含點/連字號（框架/版本）
        is_acronym = w.isupper() and 2 <= len(w) <= 8
        is_titlecase = w[0].isupper() and any(ch.islower() for ch in w[1:])
        has_sep = ("-" in w) or ("." in w) or ("_" in w)

        if is_acronym or is_titlecase or has_sep:
            candidates.append(w)

    freq = Counter(candidates)
    terms = [t for t, _ in freq.most_common(top_k)]

    # 去重但保序
    seen = set()
    uniq = []
    for t in terms:
        if t not in seen:
            uniq.append(t)
            seen.add(t)
    return uniq


def extract_internal_notes(paragraphs: list[str]) -> dict:
    prose_paras = [p for p in paragraphs if not looks_like_code(p)]
    prose_text = "\n".join(prose_paras)

    terms = extract_terms_from_text(prose_text, top_k=12)

    # 取一小段當 preview（純敘述段落）
    preview_lines = prose_paras[:8]

    return {
        "observed_terms": terms,
        "preview": "\n".join(preview_lines),
        "prose_count": len(prose_paras),
        "all_count": len(paragraphs)
    }


# -----------------------
# 故事生成（重點）
# - 禁止 keyword dump（不列清單）
# - 技術名詞以「角色」分散嵌入
# - 四段敘事：情境 → 角色與主線 → 協作流程 → 快照收斂
# -----------------------
def pick_terms_for_story(template: dict, notes: dict) -> list[str]:
    rule = template["structure"].get("required_terms_rule", {"min_terms": 3, "max_terms": 8})
    min_terms = int(rule.get("min_terms", 3))
    max_terms = int(rule.get("max_terms", 8))

    observed = notes.get("observed_terms", [])
    picked = observed[:max_terms]

    # fallback：抽不到就給占位，但不要影響敘事太多
    if len(picked) < min_terms:
        fallback = ["核心技術", "主要工具", "產出方式"]
        picked = (picked + fallback)[:max_terms]
        while len(picked) < min_terms:
            picked.append("關鍵概念")

    return picked


def generate_task01_story(template: dict, notes: dict) -> list[str]:
    picked = pick_terms_for_story(template, notes)

    # 用 3~5 個詞做「角色嵌入」，但不列成清單
    t1 = picked[0]
    t2 = picked[1] if len(picked) > 1 else "主要工具"
    t3 = picked[2] if len(picked) > 2 else "產出方式"
    t4 = picked[3] if len(picked) > 3 else None

    # 段 1：情境（你要的「先讓我懂在幹嘛」）
    p1 = (
        "你拿到這本書時，可以先不要急著鑽進細節。"
        "任務一要做的，是把它的『核心主題』翻譯成你一眼就懂的故事："
        "這本書不是在堆術語，而是在教你怎麼把一件事做成、做穩。"
    )

    # 段 2：角色與主線（技術名詞自然出現）
    p2 = (
        f"你可以把 {t1} 想成故事的主角，它決定了整本書的主線在走哪條路；"
        f"{t2} 像是主角手上的工具箱，負責把抽象想法變成可操作的步驟；"
        f"而 {t3} 則像是最後會留下的『成果形狀』，讓你能看見、能檢查、能拿去用。"
    )

    # 段 3：協作流程（你要求的「角色怎麼接力」）
    if t4:
        p3 = (
            "接下來的流程像接力：你先用主線概念把方向定住，"
            f"再用 {t2} 這類方法把事情拆成可以逐步完成的小任務；"
            f"中途會透過 {t4} 這類輔助角色處理『串接/整合/自動化』的細節，"
            f"最後把成果落在 {t3} 上，讓你可以反覆驗證、再調整、直到穩定。"
        )
    else:
        p3 = (
            "接下來的流程像接力：你先把方向定住，"
            f"再用 {t2} 這類方法把事情拆成可以逐步完成的小任務；"
            f"最後把成果落在 {t3} 上，讓你可以反覆驗證、再調整、直到穩定。"
        )

    # 段 4：快照收斂（你最重視的那句）
    snapshot_starter = template["structure"]["required_section_starter"]["snapshot"]
    p4 = (
        f"{snapshot_starter}：你腦中只要留下這張快照——"
        f"『這本書會用 {t1} 把主線定清楚，用 {t2} 把路徑拆開，"
        f"最後把成果落在 {t3} 讓你看得見、驗得到、改得動。』"
    )

    return [p1, p2, p3, p4]


# -----------------------
# 檢核（硬規格）
# - 段落數
# - 快照句
# - 技術名詞命中數（依規則）
# - 禁止 keyword dump（出現「：A、B、C」這種直接判失敗）
# -----------------------
def validate_story(template: dict, notes: dict, paragraphs: list[str]) -> None:
    if len(paragraphs) != template["structure"]["paragraphs"]:
        raise ValueError("段落數不符合模板要求")

    joined = "\n".join(paragraphs)

    # 快照句檢核
    snapshot = template["structure"]["required_section_starter"]["snapshot"]
    if snapshot not in joined:
        raise ValueError("缺少快照收斂段落的固定開頭")

    # 禁止 keyword dump：出現「：」後面緊跟多個逗號/頓號分隔
    if re.search(r"：[^。\n]{0,50}(、|,).+(、|,).+", joined):
        raise ValueError("偵測到關鍵字清單式輸出（keyword dump），不符合敘事要求")

    # 技術名詞數量檢核（通用規則）
    rule = template["structure"].get("required_terms_rule")
    if rule:
        min_terms = int(rule.get("min_terms", 3))
        max_terms = int(rule.get("max_terms", 8))
        observed = notes.get("observed_terms", [])
        hit = [t for t in observed[:max_terms] if t in joined]
        if len(hit) < min_terms:
            raise ValueError(f"技術名詞出現不足：需要至少 {min_terms} 個，實際只有 {len(hit)} 個")


# -----------------------
# 輸出 Word
# -----------------------
def write_docx(title: str, paragraphs: list[str], output_path: Path) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# -----------------------
# 主程式
# -----------------------
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", required=True)
    ap.add_argument("--source", required=True)
    ap.add_argument("--author", required=True)
    ap.add_argument("--book", required=True)
    ap.add_argument("--outdir", default="outputs")
    args = ap.parse_args()

    template = load_template(Path(args.template))
    paras = read_docx_paragraphs(Path(args.source))
    notes = extract_internal_notes(paras)

    story = generate_task01_story(template, notes)
    validate_story(template, notes, story)

    today = dt.datetime.now().strftime("%Y%m%d")
    filename = template["output"]["filename_pattern"].format(
        book=args.book, author=args.author, date=today
    )
    out_path = Path(args.outdir) / filename

    title = template.get("name", "任務一｜主題白話理解")
    write_docx(title, story, out_path)

    print("Source paragraphs:", notes.get("all_count"))
    print("Prose paragraphs:", notes.get("prose_count"))
    print("Observed terms:", notes.get("observed_terms"))
    print(f"OK: {out_path}")


if __name__ == "__main__":
    main()
